from rooms11 import proc_window, ui
from manageDB import workWithDb
import sqlite3
from PyQt5.QtWidgets import QTableWidgetItem
from messages import Warning
import docx
import datetime

import os
import getpass
def createfolder():
    try:
        USER_NAME = getpass.getuser()
        g=r"C:\Users\%s\Отчёты" % USER_NAME
        os.mkdir(g)
        return g
    except:
        return g

class rooms_work():
    def open(self):
        proc_window.show()
        fill()
        fillCombo()

    def add(self):
         workWithDb(f"INSERT INTO rooms(Field1,Field2,Field3,Field4,Field5,Field6) VALUES("
                    f"'{ui.room.text()}',"
                    f"'{ui.flour.text()}',"
                    f"'{ui.place.text()}',"
                    f"'{'Свободен'}',"
                    f"NULL,"
                    f"'{ui.Maid.currentText()}')")
         print(ui.Maid.currentText())
         ui.room.setText('')
         ui.flour.setText('')
         ui.place.setText('')
         fill()

    def delet_employee(self):  # Функиция удаления сотрудника
        returnValue = Warning()
        if returnValue == True:
            row = ui.table_proc.currentIndex().row()
            workWithDb(f"DELETE FROM rooms WHERE Field5={ui.table_proc.model().index(row, 4).data()}")
            fill()

    def create_report(self):
        g = createfolder()
        doc = docx.Document()
        # добавляем таблицу
        with sqlite3.connect('database/sanatorium.db') as db:
            cursor = db.cursor()
            for i in cursor.execute(f"SELECT COUNT(1) FROM rooms"):
                i = int(i[0]) + 1
            print(i)
            lw = [['Номер', 'Этаж', 'Мест', 'Статус' ]]
            for b in cursor.execute(
                    f"SELECT Field1,Field2,Field3,Field4 FROM rooms"):
                lw.append(b)
        paragraph = doc.add_heading('Общий отчёт', 0)
        paragraph.bold = True
        paragraph.alignment = 1

        table = doc.add_table(rows=i, cols=4)
        # применяем стиль для таблицы
        table.style = 'Table Grid'
        # заполняем таблицу данными
        for row in range(i):
            for col in range(4):
                # получаем ячейку таблицы
                cell = table.cell(row, col)
                # записываем в ячейку данные
                cell.text = str(lw[row][col])
                print(str(lw[row][col]), end='')

        # paragraph1 = doc.add_paragraph('Дата составления отчёта: ' + str(datetime.data))
        # # paragraph1.paragraph_format.space_after = Pt(12)
        # # paragraph1.paragraph_format.space_before = Pt(16)
        # paragraph1.alignment = 2
        o = str(g) + "/Отчет.docx"
        doc.save(o)

    def search(self):
        name = ui.search_entry.text()
        if name == '':
            fill()
        with sqlite3.connect('database/sanatorium.db') as db:
            cursor = db.cursor()
            #ui.table_proc.setRowCount(0)
            for surname, name, second_name, salary, phone, gg in cursor.execute(
                    f"SELECT * FROM rooms WHERE Field1 LIKE '%'||?||'%' OR Field2 LIKE '%'||?||'%' OR Field3 LIKE '%'||?||'%' OR Field4 LIKE '%'||?||'%' OR Field5 LIKE '%'||?||'%'",
                    (name, name, name, name, name)):
                row = ui.table_proc.rowCount()
                ui.table_proc.setRowCount(row + 1)
                ui.table_proc.setItem(row, 0, QTableWidgetItem(str(surname)))
                ui.table_proc.setItem(row, 1, QTableWidgetItem(str(name)))
                ui.table_proc.setItem(row, 2, QTableWidgetItem(str(second_name)))
                ui.table_proc.setItem(row, 3, QTableWidgetItem(str(salary)))
                ui.table_proc.setItem(row, 4, QTableWidgetItem(str(phone)))
                ui.table_proc.setItem(row, 5, QTableWidgetItem(str(gg)))
            db.commit()

    ui.back_button.clicked.connect(lambda : proc_window.close())
    ui.add_button.clicked.connect(add)
    ui.remove_button.clicked.connect(delet_employee)
    ui.search_button.clicked.connect(search)
    ui.report_button.clicked.connect(create_report)
    ui.table_proc.setColumnHidden(4,True)

def fillCombo():
    mas = []
    with sqlite3.connect('database/sanatorium.db') as db:
        cursor = db.cursor()
        for i in cursor.execute(f"""SELECT Field1,Field2,Field3,Field4 FROM empl"""):
            new_i = i[0] + ' ' + i[1][0] + '.' + i[2][0] + '. - ' + i[3]
            mas.append(new_i)
    ui.Maid.clear()
    ui.Maid.addItems(mas)


def fill():
    with sqlite3.connect('database/sanatorium.db') as db:
        cursor = db.cursor()
        ui.table_proc.setRowCount(0)
        for surname_name, name, second_name, GPA, phone_number,Field5 in cursor.execute(f"SELECT Field1,Field2,Field3,Field4,Field5,Field6 FROM rooms"):
            row = ui.table_proc.rowCount()
            ui.table_proc.setRowCount(row + 1)
            ui.table_proc.setItem(row, 0, QTableWidgetItem(str(surname_name)))
            ui.table_proc.setItem(row, 1, QTableWidgetItem(str(name)))
            ui.table_proc.setItem(row, 2, QTableWidgetItem(str(second_name)))
            ui.table_proc.setItem(row, 3, QTableWidgetItem(str(GPA)))
            ui.table_proc.setItem(row, 4, QTableWidgetItem(str(phone_number)))
            ui.table_proc.setItem(row, 5, QTableWidgetItem(str(Field5)))
        db.commit()

#Создаем папку для отчета

# import sys
# app = QtWidgets.QApplication(sys.argv)
# proc_window = QtWidgets.QWidget()
# ui = Ui_proc_window()
# ui.setupUi(proc_window)